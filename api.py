import warnings

warnings.filterwarnings("ignore")
import json
import logging
import random
import time
import uvicorn
from fastapi import FastAPI, status, Depends, Request, Body
from fastapi.responses import JSONResponse, StreamingResponse
from sqlalchemy import create_engine, text, select, insert, delete, func, MetaData, Table, Column, Boolean, Text
from sqlalchemy.exc import IntegrityError, OperationalError
from sqlalchemy.orm import Session, sessionmaker, scoped_session
from sqlalchemy.pool import QueuePool
from datetime import datetime, timedelta
import jieba.analyse
from wordcloud import WordCloud, STOPWORDS
import io
import base64
import threading
from utils.table_models import DailyData, PushedDailyData, Data, HumanLLM, Sen, TempHumanLLM, TempSen, Users, LLMRank
from urllib.parse import urlencode
from apscheduler.schedulers.background import BackgroundScheduler
from apscheduler.triggers.interval import IntervalTrigger
from uuid import uuid4
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Cm

jieba.setLogLevel(logging.ERROR)

app = FastAPI()
schedular = BackgroundScheduler()
report_data_store = {}
cached_item_store = {}
cache_expiry_times = {}

llm_labels = ["kimi", "通义", "文心一言", "智谱", "ChatGLM3-6B", "QWen1.5-7B", "QWen1.5-14B", "Baichuan2-7B",
              "Baichuan2-13B", "ChatGPT", "Llama2-7B", "Llama2-13B", "Llama3-8B", "Mistral v0.2 7B"]
type_list = ['境外新闻', '社交媒体', '消息应用', '问答社区']
topic_list = ['领导人', '南海', '台湾', '新疆', '西藏']

stopwords = set()
with open('./utils/stopwords.txt', 'r', encoding='utf-8') as f:
    for line in f.readlines():
        stopwords.add(line.strip())

with open('./utils/topic_mapping.json', 'r', encoding='utf-8') as f:
    topic_mapping_dict = json.load(f)


DATABASE_DATA = 'sqlite:///./database/data.db'
DATABASE_DAILY_DATA = 'sqlite:///./database/daily_data.db'
DATABASE_SEN = 'sqlite:///./database/sen.db'
DATABASE_HUMAN_LLM = 'sqlite:///./database/human_llm.db'
DATABASE_USERS = 'sqlite:///./database/users.db'
DATABASE_TEMP_SEN = 'sqlite:///./database/temp_sen.db'
DATABASE_TEMP_HUMAN_LLM = 'sqlite:///./database/temp_human_llm.db'
DATABASE_LLM_RANK = 'sqlite:///./database/llm_rank.db'

engine_data = create_engine(DATABASE_DATA, poolclass=QueuePool, pool_size=10, max_overflow=20)
engine_daily_data = create_engine(DATABASE_DAILY_DATA, poolclass=QueuePool, pool_size=10, max_overflow=20)
engine_sen = create_engine(DATABASE_SEN, poolclass=QueuePool, pool_size=10, max_overflow=20)
engine_human_llm = create_engine(DATABASE_HUMAN_LLM, poolclass=QueuePool, pool_size=10, max_overflow=20)
engine_users = create_engine(DATABASE_USERS, poolclass=QueuePool, pool_size=5, max_overflow=10)
engine_temp_sen = create_engine(DATABASE_TEMP_SEN, poolclass=QueuePool, pool_size=10, max_overflow=20)
engine_temp_human_llm = create_engine(DATABASE_TEMP_HUMAN_LLM, poolclass=QueuePool, pool_size=10, max_overflow=20)
engine_llm_rank = create_engine(DATABASE_LLM_RANK, poolclass=QueuePool, pool_size=5, max_overflow=10)

SessionLocalData = scoped_session(sessionmaker(bind=engine_data, autoflush=False, autocommit=False))
SessionLocalDailyData = scoped_session(sessionmaker(bind=engine_daily_data, autoflush=False, autocommit=False))
SessionLocalSen = scoped_session(sessionmaker(bind=engine_sen, autoflush=False, autocommit=False))
SessionLocalHumanLLM = scoped_session(sessionmaker(bind=engine_human_llm, autoflush=False, autocommit=False))
SessionLocalUsers = scoped_session(sessionmaker(bind=engine_users, autoflush=False, autocommit=False))
SessionLocalTempSen = scoped_session(sessionmaker(bind=engine_temp_sen, autoflush=False, autocommit=False))
SessionLocalTempHumanLLM = scoped_session(sessionmaker(bind=engine_temp_human_llm, autoflush=False, autocommit=False))
SessionLocalLLMRank = scoped_session(sessionmaker(bind=engine_llm_rank, autoflush=False, autocommit=False))


def get_db(session_local):
    def db():
        db = session_local()
        try:
            yield db
        finally:
            db.close()

    return db


metadata = MetaData()


def topic_map(content):
    topics = []
    for i in topic_list:
        for kw in topic_mapping_dict[i]:
            if kw in content:
                topics.append(i)
                break

    return topics


def add_column_if_not_exists(engine, table_name, column, default=None):
    metadata.bind = engine
    table = Table(table_name, metadata, autoload_with=engine)

    if column.name not in table.c:
        with engine.connect() as conn:
            try:
                if default is not None:
                    conn.execute(
                        text(f'ALTER TABLE {table_name} ADD COLUMN {column.name} {column.type} DEFAULT {default}'))
                else:
                    conn.execute(text(f'ALTER TABLE {table_name} ADD COLUMN {column.name} {column.type}'))
                print(f'Added column {column.name} to {table_name}')
            except OperationalError as e:
                print(f'Error adding column {column.name} to {table_name}: {e}')


def init_db():
    Data.__table__.create(engine_data, checkfirst=True)
    DailyData.__table__.create(engine_daily_data, checkfirst=True)
    PushedDailyData.__table__.create(engine_daily_data, checkfirst=True)
    Sen.__table__.create(engine_sen, checkfirst=True)
    HumanLLM.__table__.create(engine_human_llm, checkfirst=True)
    Users.__table__.create(engine_users, checkfirst=True)
    TempSen.__table__.create(engine_temp_sen, checkfirst=True)
    TempHumanLLM.__table__.create(engine_temp_human_llm, checkfirst=True)
    LLMRank.__table__.create(engine_llm_rank, checkfirst=True)

    # add_column_if_not_exists(engine_sen, 'sen', Column('push_time', Text),
    #                          "'" + datetime.now().strftime('%Y-%m-%d %H:%M:%S') + "'")
    # add_column_if_not_exists(engine_human_llm, 'human_llm', Column('push_time', Text),
    #                          "'" + datetime.now().strftime('%Y-%m-%d %H:%M:%S') + "'")
    # add_column_if_not_exists(engine_temp_sen, 'temp_sen', Column('push_time', Text),
    #                          "'" + datetime.now().strftime('%Y-%m-%d %H:%M:%S') + "'")
    # add_column_if_not_exists(engine_temp_human_llm, 'temp_human_llm', Column('push_time', Text),
    #                          "'" + datetime.now().strftime('%Y-%m-%d %H:%M:%S') + "'")
    # add_column_if_not_exists(engine_data, 'data', Column('push_time', Text),
    #                          "'" + datetime.now().strftime('%Y-%m-%d %H:%M:%S') + "'")
    # add_column_if_not_exists(engine_daily_data, 'daily_data', Column('submitted', Text), default='FALSE')
    # add_column_if_not_exists(engine_human_llm, 'human_llm', Column('submitted_time', Text),
    #                          default="'" + datetime.now().strftime('%Y-%m-%d %H:%M:%S') + "'")
    # add_column_if_not_exists(engine_sen, 'sen', Column('submitted_time', Text),
    #                          default="'" + datetime.now().strftime('%Y-%m-%d %H:%M:%S') + "'")

    daily_daba_db = SessionLocalDailyData()
    daily_daba_db.execute(delete(DailyData).filter(DailyData.submitted == True, func.strftime(DailyData.push_time,
                                                                                              '%Y-%m-%d') != datetime.now().strftime(
        '%Y-%m-%d')))
    daily_daba_db.commit()
    daily_daba_db.close()


def key_builder(**kwargs) -> str:
    return '_'.join([f"{key}:{value or 'none'}" for key, value in sorted(kwargs.items())])


def set_cache(key: str, data: dict, expiry_seconds: int):
    cached_item_store[key] = data
    cache_expiry_times[key] = datetime.now() + timedelta(seconds=expiry_seconds)


@app.post('/push_data')
def push_data(type: str = Body(...), date_time: str = Body(...),
              content: str = Body(...), url: str = Body(None),
              daily_data_db: Session = Depends(get_db(SessionLocalDailyData))) -> JSONResponse:
    if type not in type_list:
        return JSONResponse(status_code=status.HTTP_400_BAD_REQUEST, content={'message': '输入类型错误'})
    try:
        topic = topic_map(content)

        result = daily_data_db.execute(
            insert(DailyData).values(type=type, time=datetime.strptime(date_time, '%Y-%m-%d %H:%M:%S'),
                                     push_time=datetime.now().strftime('%Y-%m-%d %H:%M:%S'), content=content,
                                     sensitive=False, sensitive_score=0., is_bot=False, is_bot_score=0.,
                                     model_judgment=None, url=url, topic=' '.join(topic) if len(topic) else None))
        data_id = result.lastrowid
        daily_data_db.execute(insert(PushedDailyData).values(data_id=data_id))
        daily_data_db.commit()

        row_count = daily_data_db.query(DailyData.ID).count()

        return JSONResponse(content={'daily_count': row_count})
    except:
        return JSONResponse(status_code=status.HTTP_400_BAD_REQUEST, content={'message': 'Bad Request'})


@app.get('/daily_data')
def read_daily_data(page: int = 1, page_size: int = 10,
                    db: Session = Depends(get_db(SessionLocalDailyData))) -> JSONResponse:
    try:
        results = db.query(DailyData).limit(page_size).offset((page - 1) * page_size).all()

        rows = []
        for row in results:
            rows.append([row.ID, row.type, row.time, row.content, row.sensitive, row.sensitive_score, row.is_bot,
                         row.is_bot_score, row.model_judgment, row.url, row.topic])

        return JSONResponse(content={'daily_data': rows})
    except:
        return JSONResponse(status_code=status.HTTP_400_BAD_REQUEST, content={'message': 'Bad Request'})


@app.get('/data')
def read_data(page: int = 1, page_size: int = 10,
              db: Session = Depends(get_db(SessionLocalData))) -> JSONResponse:
    try:
        results = db.query(Data).limit(page_size).offset((page - 1) * page_size).all()

        rows = []
        for row in results:
            rows.append([row.ID, row.type, row.time, row.content, row.sensitive, row.sensitive_score, row.is_bot,
                         row.is_bot_score, row.model_judgment, row.url, row.topic])

        return JSONResponse(content={'data': rows})
    except:
        return JSONResponse(status_code=status.HTTP_400_BAD_REQUEST, content={'message': 'Bad Request'})


@app.get('/daily_data_count')
def daily_data_count(db: Session = Depends(get_db(SessionLocalDailyData))) -> JSONResponse:
    try:
        count = db.query(DailyData.ID).count()

        return JSONResponse(content={'daily_data_count': count})
    except Exception as e:
        return JSONResponse(status_code=status.HTTP_500_INTERNAL_SERVER_ERROR, content={'message': str(e)})


@app.get('/data_count')
def data_count(db: Session = Depends(get_db(SessionLocalData))) -> JSONResponse:
    try:
        count = db.query(Data.ID).count()

        return JSONResponse(content={'data_count': count})
    except Exception as e:
        return JSONResponse(status_code=status.HTTP_500_INTERNAL_SERVER_ERROR, content={'message': str(e)})


# @app.get('/sensitive')
# def sensitive(content: str) -> JSONResponse:
#     try:
#         is_sensitive, score = predict_sensitive(content)
#         response = {
#             'is_sensitive': is_sensitive,
#             'score': score
#         }
#         return JSONResponse(content=response)
#     except:
#         return JSONResponse(status_code=status.HTTP_400_BAD_REQUEST, content={'message': 'Bad Request'})


# @app.get('/llm_det')
# def llm_det(content: str) -> JSONResponse:
#     try:
#         response = predict_llm(content)
#         return JSONResponse(content=response)
#     except:
#         return JSONResponse(status_code=status.HTTP_400_BAD_REQUEST, content={'message': 'Bad Request'})


@app.get('/search/daily_data')
def search_daily_data(page: int = 1, page_size: int = 10, type: str = None, sensitive: bool = None,
                      is_bot: bool = None, model_judgment: str = None,
                      content_keyword: str = None,
                      topic: str = None,
                      db: Session = Depends(get_db(SessionLocalDailyData))) -> JSONResponse:
    try:
        filters = []
        filters.append(DailyData.submitted == False)
        if type:
            filters.append(DailyData.type == type)
        if sensitive is not None:
            filters.append(DailyData.sensitive == sensitive)
        if is_bot is not None:
            filters.append(DailyData.is_bot == is_bot)
        if model_judgment:
            filters.append(DailyData.model_judgment == model_judgment)
        if content_keyword:
            filters.append(DailyData.content.like(f'%{content_keyword}%'))
        if topic:
            filters.append(DailyData.topic.like(f'%{topic}%'))

        query = db.query(DailyData).filter(*filters)
        total_count = query.count()

        results = query.limit(page_size).offset((page - 1) * page_size).all()

        rows = []
        for row in results:
            rows.append(
                [row.ID, row.type, row.time, row.content, row.sensitive, round(row.sensitive_score, 2), row.is_bot,
                 round(row.is_bot_score, 2), row.model_judgment, row.url, row.topic])

        return JSONResponse(content={'total_count': total_count, 'data': rows})
    except:
        return JSONResponse(status_code=status.HTTP_400_BAD_REQUEST, content={'message': 'Bad Request'})


@app.get('/search/human_llm')
def search_human_llm(page: int = 1, page_size: int = 10, type: str = None, is_bot: bool = None,
                     model_judgment: str = None,
                     content_keyword: str = None,
                     topic: str = None,
                     db: Session = Depends(get_db(SessionLocalHumanLLM))) -> JSONResponse:
    try:
        filters = []
        if type:
            filters.append(HumanLLM.type == type)
        if is_bot is not None:
            filters.append(HumanLLM.is_bot == is_bot)
        if model_judgment:
            filters.append(HumanLLM.model_judgment == model_judgment)
        if content_keyword:
            filters.append(HumanLLM.content.like(f'%{content_keyword}%'))
        if topic:
            filters.append(HumanLLM.topic.split(f'%{topic}%'))

        query = db.query(HumanLLM).filter(*filters)
        total_count = query.count()

        results = query.limit(page_size).offset((page - 1) * page_size).all()

        rows = []
        for row in results:
            rows.append({'ID': row.ID, 'type': row.type, 'time': row.time, 'content': row.content, 'is_bot': row.is_bot,
                         'is_bot_score': round(row.is_bot_score, 2), 'model_judgment': row.model_judgment,
                         'url': row.url,
                         'topic': row.topic, 'auditor': row.auditor})

        return JSONResponse(content={'total_count': total_count, 'data': rows})
    except Exception as e:
        return JSONResponse(status_code=status.HTTP_400_BAD_REQUEST, content={'message': str(e)})


@app.get('/search/sen')
def search_sen(page: int = 1, page_size: int = 10, type: str = None, sensitive: bool = None,
               content_keyword: str = None,
               topic: str = None,
               db: Session = Depends(get_db(SessionLocalSen))) -> JSONResponse:
    try:
        filters = []
        if type:
            filters.append(Sen.type == type)
        if sensitive is not None:
            filters.append(Sen.sensitive == sensitive)
        if content_keyword:
            filters.append(Sen.content.like(f'%{content_keyword}%'))
        if topic:
            filters.append(Sen.topic.like(f'%{topic}%'))

        query = db.query(Sen).filter(*filters)
        total_count = query.count()

        results = query.limit(page_size).offset((page - 1) * page_size).all()

        rows = []
        for row in results:
            rows.append(
                {'ID': row.ID, 'type': row.type, 'time': row.time, 'content': row.content, 'sensitive': row.sensitive,
                 'sensitive_score': round(row.sensitive_score, 2), 'url': row.url, 'topic': row.topic,
                 'auditor': row.auditor})

        return JSONResponse(content={'total_count': total_count, 'data': rows})
    except Exception as e:
        return JSONResponse(status_code=status.HTTP_400_BAD_REQUEST, content={'message': str(e)})


@app.get('/llm_rank')
def llm_rank(start_date: str, end_date: str,
             human_llm_db: Session = Depends(get_db(SessionLocalHumanLLM)),
             llm_rank_db: Session = Depends(get_db(SessionLocalLLMRank))) -> JSONResponse:
    try:
        start_date = datetime.strptime(start_date, '%Y-%m-%d')
        end_date = datetime.strptime(end_date, '%Y-%m-%d')

        results = human_llm_db.execute(
            select(HumanLLM.type, HumanLLM.model_judgment).where(HumanLLM.time.between(start_date, end_date),
                                                                 HumanLLM.is_bot == True,
                                                                 HumanLLM.model_judgment != None)).fetchall()

        n = 0
        model_count = {}
        for type in type_list:
            model_count[type] = {}
            for model in llm_labels:
                model_count[type][model] = 0.
        for row in results:
            model_count[row.type][row.model_judgment] += 1.
            n += 1

        if n == 0:
            n = 1

        for type in type_list:
            for model in llm_labels:
                model_count[type][model] = model_count[type][model] * 100 / n

                record = llm_rank_db.query(LLMRank).filter(LLMRank.second_level == type, LLMRank.model == model).first()
                record.score = model_count[type][model]
                llm_rank_db.commit()

        records = llm_rank_db.query(LLMRank.first_level, LLMRank.second_level, LLMRank.metric, LLMRank.model,
                                    LLMRank.score).all()

        results = {}
        for model in llm_labels:
            results[model] = {}

        for record in records:
            first_level, second_level, metric, model, score = record.first_level, record.second_level, record.metric, record.model, record.score
            if model == 'GPT4':
                continue
            if first_level not in results[model]:
                results[model][first_level] = {}
            results[model][first_level][second_level] = score

        with open('./utils/LLM.json', 'r', encoding='utf-8') as f:
            llm_source_map = json.load(f)

        new_results = []
        for model in llm_labels:
            temp_dict = {'model': model, 'source': llm_source_map[model]}
            for first_level in results[model]:
                sum_score = 0.
                n = 0
                for second_level in results[model][first_level]:
                    score = results[model][first_level][second_level]
                    temp_dict[second_level] = round(score, 2) if score else None
                    if score:
                        sum_score += score
                        n += 1
                temp_dict[first_level] = round(sum_score / n, 2) if n else None
            new_results.append(temp_dict)

        return JSONResponse(content={'llm_rank': new_results})

    except Exception as e:
        return JSONResponse(status_code=status.HTTP_500_INTERNAL_SERVER_ERROR, content={'message': str(e)})


def get_cached_paint_data(activate: str, start_date: str = None, end_date: str = None):
    cache_key = key_builder(activate=activate, start_date=start_date, end_date=end_date)
    cached_response = cached_item_store.get(cache_key, None)
    if cached_response:
        return cached_response

    response = {}
    activate = activate.split('&')
    today = datetime.now().strftime('%Y-%m-%d')
    parameters = {}
    added_query = ''
    if start_date and end_date:
        added_query += ' AND time BETWEEN :sd AND :ed'
        parameters['sd'] = datetime.strptime(start_date, '%Y-%m-%d')
        parameters['ed'] = datetime.strptime(end_date, '%Y-%m-%d')
    elif start_date:
        added_query += ' AND time >= :sd'
        parameters['sd'] = datetime.strptime(start_date, '%Y-%m-%d')
    elif end_date:
        added_query += ' AND time <= :ed'
        parameters['ed'] = datetime.strptime(end_date, '%Y-%m-%d')

    if set(activate) & {'human_llm_count', 'human_llm_type_distribution', 'data_human_llm', 'today_human_llm_count',
                        'topic_human_llm', 'human_llm_time_distribution', 'human_llm_is_bot_type'}:
        human_llm_db = SessionLocalHumanLLM()
        if 'human_llm_count' in activate:  # 两个语料库总量
            response['human_llm_count'] = human_llm_db.query(HumanLLM.ID).count()
        if 'human_llm_type_distribution' in activate:  # 人机语料库语料来源分布
            results = human_llm_db.query(HumanLLM.type, func.count(HumanLLM.ID)).filter(
                HumanLLM.type != None).group_by(
                HumanLLM.type).all()
            if results is not None:
                response['human_llm_type_distribution'] = {row.type: row[1] for row in results}

        base_query = '''
        SELECT strftime('%Y-%m-%d', time), type, {column}, COUNT(*)
        FROM {database} WHERE time IS NOT NULL AND type IS NOT NULL
        ''' + added_query
        if 'data_human_llm' in activate:  # 大模型生成数量趋势
            results = human_llm_db.execute(text(
                base_query.format(column='is_bot', database='human_llm') + " GROUP BY time, type, is_bot"),
                parameters).fetchall()
            if results is not None:
                data_human_llm = {}
                for row in results:
                    if row[0] not in data_human_llm:
                        data_human_llm[row[0]] = {}
                        for type in type_list:
                            data_human_llm[row[0]][type] = {1: 0, 0: 0}
                    data_human_llm[row[0]][row[1]][row[2]] = row[3]
                response['data_human_llm'] = dict(sorted(data_human_llm.items(), key=lambda item: item[0]))

        if 'today_human_llm_count' in activate:  # 今日敏感数据库新增
            response['today_human_llm_count'] = human_llm_db.query(HumanLLM.ID).filter(
                func.strftime('%Y-%m-%d', HumanLLM.submitted_time) == today).count()

        if 'topic_human_llm' in activate:  # 话题-人机生成量
            results = human_llm_db.query(HumanLLM.topic, HumanLLM.is_bot, func.count(HumanLLM.ID)).filter(
                HumanLLM.topic != None).group_by(HumanLLM.topic, HumanLLM.is_bot).all()
            if results is not None:
                topic_human_llm = {}
                for topic in topic_list:
                    topic_human_llm[topic] = {1: 0, 0: 0}
                for row in results:
                    row_topics = row.topic.split(' ')
                    for row_topic in row_topics:
                        topic_human_llm[row_topic][row.is_bot] += row[2]
                response['topic_human_llm'] = topic_human_llm

        base_query = '''
        SELECT strftime('%Y-%m-%d', time), type, COUNT(*)
        FROM {database} WHERE time IS NOT NULL AND type IS NOT NULL
        ''' + added_query

        if 'human_llm_time_distribution' in activate:
            results = human_llm_db.execute(text(
                base_query.format(database='human_llm') + " GROUP BY time, type"),
                parameters).fetchall()
            if results is not None:
                human_llm_time_distribution = {}
                for row in results:
                    if row[0] not in human_llm_time_distribution:
                        human_llm_time_distribution[row[0]] = {}
                        for type in type_list:
                            human_llm_time_distribution[row[0]][type] = 0
                    human_llm_time_distribution[row[0]][row[1]] = row[2]

                response['human_llm_time_distribution'] = dict(
                    sorted(human_llm_time_distribution.items(), key=lambda item: item[0]))

        if 'human_llm_is_bot_type' in activate:
            results = human_llm_db.query(HumanLLM.type, func.count(HumanLLM.ID)).filter(
                HumanLLM.type != None, HumanLLM.is_bot == True).group_by(
                HumanLLM.type).all()
            if results is not None:
                response['human_llm_type_distribution'] = {row.type: row[1] for row in results}

        human_llm_db.close()

    if set(activate) & {'sen_count', 'sen_type_distribution', 'data_sensitive', 'today_sen_count', 'topic_sen',
                        'sen_time_distribution', 'sen_sensitive_type'}:
        sen_db = SessionLocalSen()
        if 'sen_count' in activate:  # 两个语料库总量
            response['sen_count'] = sen_db.query(Sen.ID).count()
        if 'sen_type_distribution' in activate:  # 敏感数据库语料来源分布
            results = sen_db.query(Sen.type, func.count(Sen.ID)).filter(Sen.type != None).group_by(Sen.type).all()
            if results is not None:
                response['sen_type_distribution'] = {row.type: row[1] for row in results}

        base_query = '''
        SELECT strftime('%Y-%m-%d', time), type, {column}, COUNT(*)
        FROM {database} WHERE time IS NOT NULL AND type IS NOT NULL
        ''' + added_query
        if 'data_sensitive' in activate:  # 敏感数量趋势
            results = sen_db.execute(text(
                base_query.format(column='sensitive', database='sen') + " GROUP BY time, type, sensitive"),
                parameters).fetchall()
            if results is not None:
                data_sensitive = {}
                for row in results:
                    if row[0] not in data_sensitive:
                        data_sensitive[row[0]] = {}
                        for type in type_list:
                            data_sensitive[row[0]][type] = {1: 0, 0: 0}
                    data_sensitive[row[0]][row[1]][row[2]] = row[3]
                response['data_sensitive'] = dict(sorted(data_sensitive.items(), key=lambda item: item[0]))

        if 'today_sen_count' in activate:  # 今日敏感数据库新增
            response['today_sen_count'] = sen_db.query(Sen.ID).filter(
                func.strftime('%Y-%m-%d', Sen.submitted_time) == today).count()

        if 'topic_sen' in activate:  # 话题-敏感信息量
            results = sen_db.query(Sen.topic, Sen.sensitive, func.count(Sen.ID)).filter(
                Sen.topic != None).group_by(Sen.topic, Sen.sensitive).all()
            if results is not None:
                topic_sen = {}
                for topic in topic_list:
                    topic_sen[topic] = {1: 0, 0: 0}
                for row in results:
                    row_topics = row.topic.split(' ')
                    for row_topic in row_topics:
                        topic_sen[row_topic][row.sensitive] += row[2]
                response['topic_sen'] = topic_sen

        base_query = '''
        SELECT strftime('%Y-%m-%d', time), type, COUNT(*)
        FROM {database} WHERE time IS NOT NULL AND type IS NOT NULL
        ''' + added_query
        if 'sen_time_distribution' in activate:
            results = sen_db.execute(text(
                base_query.format(database='sen') + " GROUP BY time, type"),
                parameters).fetchall()
            if results is not None:
                sen_time_distribution = {}
                for row in results:
                    if row[0] not in sen_time_distribution:
                        sen_time_distribution[row[0]] = {}
                        for type in type_list:
                            sen_time_distribution[row[0]][type] = 0
                    sen_time_distribution[row[0]][row[1]] = row[2]

                response['sen_time_distribution'] = dict(
                    sorted(sen_time_distribution.items(), key=lambda item: item[0]))

        if 'sen_sensitive_type' in activate:
            results = sen_db.query(Sen.type, func.count(Sen.ID)).filter(
                Sen.type != None, Sen.sensitive == True).group_by(
                Sen.type).all()
            if results is not None:
                response['sen_sensitive_type'] = {row.type: row[1] for row in results}

        sen_db.close()

    if set(activate) & {'data_time_distribution', 'topic_count'}:
        data_db = SessionLocalData()
        base_query = '''
        SELECT strftime('%Y-%m-%d', time), type, COUNT(*)
        FROM {database} WHERE time IS NOT NULL AND type IS NOT NULL
        ''' + added_query

        if 'data_time_distribution' in activate:  # 历史数据量统计（四通道柱状图）
            results = data_db.execute(text(
                base_query.format(database='data') + " GROUP BY time, type"),
                parameters).fetchall()
            data_time_distribution = {}
            for row in results:
                if row[0] not in data_time_distribution:
                    data_time_distribution[row[0]] = {}
                    for type in type_list:
                        data_time_distribution[row[0]][type] = 0
                data_time_distribution[row[0]][row[1]] = row[2]

            response['data_time_distribution'] = dict(sorted(data_time_distribution.items(), key=lambda item: item[0]))

        if 'topic_count' in activate:  # 话题信息总量
            results = data_db.query(Data.topic, func.count(Data.ID)).filter(
                Data.topic != None).group_by(Data.topic).all()
            if results is not None:
                topic_count = {}
                for topic in topic_list:
                    topic_count[topic] = 0
                for row in results:
                    row_topics = row.topic.split(' ')
                    for row_topic in row_topics:
                        topic_count[row_topic] += row[1]

                response['topic_count'] = topic_count

        data_db.close()

    if set(activate) & {'today_data_count', 'today_topic_type', 'today_topic_count', 'today_topic_sen',
                        'today_topic_human_llm'}:
        daily_data_db = SessionLocalDailyData()
        if 'today_data_count' in activate:  # 今日到来
            response['today_data_count'] = daily_data_db.query(DailyData.ID).filter(
                func.strftime('%Y-%m-%d', DailyData.push_time) == today).count()

        if 'today_topic_type' in activate:  # 今日 话题-来源量
            results = daily_data_db.query(DailyData.topic, DailyData.type, func.count(DailyData.ID)).filter(
                DailyData.topic != None, DailyData.type != None,
                func.strftime('%Y-%m-%d', DailyData.push_time) == today).group_by(
                DailyData.topic, DailyData.type).all()
            if results is not None:
                today_topic_type = {}
                for topic in topic_list:
                    today_topic_type[topic] = {}
                    for type in type_list:
                        today_topic_type[topic][type] = 0
                    today_topic_type[topic][None] = 0
                for row in results:
                    row_topics = row.topic.split(' ')
                    for row_topic in row_topics:
                        today_topic_type[row_topic][row.type] += row[2]

                response['today_topic_type'] = today_topic_type

        if 'today_topic_count' in activate:  # 今日 话题信息量
            results = daily_data_db.query(DailyData.topic, func.count(DailyData.ID)).filter(
                DailyData.topic != None, func.strftime('%Y-%m-%d', DailyData.push_time) == today).group_by(
                DailyData.topic).all()
            if results is not None:
                today_topic_count = {}
                for topic in topic_list:
                    today_topic_count[topic] = 0
                for row in results:
                    row_topics = row.topic.split(' ')
                    for row_topic in row_topics:
                        today_topic_count[row_topic] += row[1]

                response['today_topic_count'] = today_topic_count

        if 'today_topic_sen' in activate:  # 今日 话题-敏感信息量
            results = daily_data_db.query(DailyData.topic, DailyData.sensitive, func.count(DailyData.ID)).filter(
                DailyData.topic != None, func.strftime('%Y-%m-%d', DailyData.push_time) == today).group_by(
                DailyData.topic, DailyData.sensitive).all()
            if results is not None:
                today_topic_sen = {}
                for topic in topic_list:
                    today_topic_sen[topic] = {1: 0, 0: 0}
                for row in results:
                    row_topics = row.topic.split(' ')
                    for row_topic in row_topics:
                        today_topic_sen[row_topic][row.sensitive] += row[2]

                response['today_topic_sen'] = today_topic_sen

        if 'today_topic_human_llm' in activate:  # 今日 话题-人机生成量
            results = daily_data_db.query(DailyData.topic, DailyData.is_bot, func.count(DailyData.ID)).filter(
                DailyData.topic != None, func.strftime('%Y-%m-%d', DailyData.push_time) == today).group_by(
                DailyData.topic, DailyData.is_bot).all()
            if results is not None:
                today_topic_human_llm = {}
                for topic in topic_list:
                    today_topic_human_llm[topic] = {1: 0, 0: 0}
                for row in results:
                    row_topics = row.topic.split(' ')
                    for row_topic in row_topics:
                        today_topic_human_llm[row_topic][row.is_bot] += row[2]

                response['today_topic_human_llm'] = today_topic_human_llm

        daily_data_db.close()

    if response:
        set_cache(cache_key, response, expiry_seconds=30)

    return response


@app.get('/paint_data')
def get_paint_data(activate: str, start_date: str = None, end_date: str = None) -> JSONResponse:
    try:
        response = get_cached_paint_data(activate, start_date, end_date)

        # if not response:
        #     return JSONResponse(status_code=status.HTTP_404_NOT_FOUND, content={'message': 'No Record'})

        return JSONResponse(content=response)
    except Exception as e:
        import traceback
        error = traceback.format_exc()
        print(error)
        return JSONResponse(status_code=status.HTTP_500_INTERNAL_SERVER_ERROR, content={'message': str(e)})


# @app.put('/update/daily_data/{ID}')
# def update_daily_data(ID: int, is_sensitive: bool, is_bot: bool,
#                       daily_data_db: Session = Depends(get_db(SessionLocalDailyData))):
#     try:
#         record = daily_data_db.query(DailyData).filter(DailyData.ID == ID).first()
#         if record is not None:
#             record.sensitive = is_sensitive
#             record.sensitive_score = float(is_sensitive)
#             record.is_bot = is_bot
#             record.is_bot_score = float(is_bot)
#             if is_bot is False:
#                 record.model_judgment = None
#
#             daily_data_db.commit()
#
#         return JSONResponse(status_code=status.HTTP_200_OK, content={'message': 'Updated Successfully'})
#     except Exception as e:
#         return JSONResponse(status_code=status.HTTP_500_INTERNAL_SERVER_ERROR, content=str(e))


@app.put('/submit')
def submit(ID: str, is_sensitive: str, is_bot: str, username: str,
           data_db: Session = Depends(get_db(SessionLocalData)),
           daily_data_db: Session = Depends(get_db(SessionLocalDailyData)),
           temp_sen_db: Session = Depends(get_db(SessionLocalTempSen)),
           temp_human_llm_db: Session = Depends(get_db(SessionLocalTempHumanLLM))) -> JSONResponse:
    try:
        IDs = list(map(int, ID.split('&')))
        is_sensitive_s = list(map(int, is_sensitive.split('&')))
        is_bot_s = list(map(int, is_bot.split('&')))
        username = username.split('&')[0]
    except Exception as e:
        return JSONResponse(status_code=status.HTTP_400_BAD_REQUEST, content={'message': str(e)})
    try:
        for ID_, is_sensitive_, is_bot_ in zip(IDs, is_sensitive_s, is_bot_s):
            record = daily_data_db.query(DailyData).filter(DailyData.ID == ID_, DailyData.submitted == False).first()
            if not record:
                continue
            new_sensitive_score = float(is_sensitive_) if is_sensitive_ != record.sensitive else record.sensitive_score
            new_is_bot_score = float(is_bot_) if is_bot_ != record.is_bot else record.is_bot_score
            new_model_judgment = record.model_judgment if is_bot_ else None
            uuid = uuid4().hex

            temp_sen_db.execute(insert(TempSen).values(type=record.type,
                                                       time=record.time,
                                                       push_time=record.push_time,
                                                       content=record.content,
                                                       sensitive=is_sensitive_,
                                                       sensitive_score=new_sensitive_score,
                                                       url=record.url,
                                                       topic=record.topic,
                                                       username=username,
                                                       uuid=uuid
                                                       ))
            temp_sen_db.commit()

            temp_human_llm_db.execute(insert(TempHumanLLM).values(type=record.type,
                                                                  time=record.time,
                                                                  push_time=record.push_time,
                                                                  content=record.content,
                                                                  is_bot=is_bot_,
                                                                  is_bot_score=new_is_bot_score,
                                                                  model_judgment=new_model_judgment,
                                                                  url=record.url,
                                                                  topic=record.topic,
                                                                  username=username,
                                                                  uuid=uuid
                                                                  ))
            temp_human_llm_db.commit()

            data_db.execute(insert(Data).values(type=record.type,
                                                time=record.time,
                                                push_time=record.time,
                                                content=record.content,
                                                sensitive=is_sensitive_,
                                                sensitive_score=new_sensitive_score,
                                                is_bot=is_bot_,
                                                is_bot_score=new_is_bot_score,
                                                model_judgment=new_model_judgment,
                                                url=record.url,
                                                topic=record.topic,
                                                uuid=uuid
                                                ))
            data_db.commit()

            record.submitted = True
            daily_data_db.commit()

        return JSONResponse(content={'message': 'Update Submitted'})

    except Exception as e:
        return JSONResponse(status_code=status.HTTP_500_INTERNAL_SERVER_ERROR, content={'message': str(e)})


@app.put('/submit_all')
def submit_all(username: str, type: str = None, sensitive: bool = None, is_bot: bool = None, model_judgment: str = None,
               content_keyword: str = None, topic: str = None,
               data_db: Session = Depends(get_db(SessionLocalData)),
               daily_data_db: Session = Depends(get_db(SessionLocalDailyData)),
               temp_sen_db: Session = Depends(get_db(SessionLocalTempSen)),
               temp_human_llm_db: Session = Depends(get_db(SessionLocalTempHumanLLM))) -> JSONResponse:
    try:
        filters = []
        filters.append(DailyData.submitted == False)
        if type:
            filters.append(DailyData.type == type)
        if sensitive is not None:
            filters.append(DailyData.sensitive == sensitive)
        if is_bot is not None:
            filters.append(DailyData.is_bot == is_bot)
        if model_judgment:
            filters.append(DailyData.model_judgment == model_judgment)
        if content_keyword:
            filters.append(DailyData.content.like(f'%{content_keyword}%'))
        if topic:
            filters.append(DailyData.topic.like(f'%{topic}%'))

        query = daily_data_db.query(DailyData).filter(*filters)
        total_count = query.count()

        batch_size = 1000
        total_batches = (total_count // batch_size) + 1

        for batch in range(total_batches):
            results = query.limit(batch_size).offset(batch * batch_size).all()
            for record in results:
                uuid = uuid4().hex
                temp_sen_db.execute(insert(TempSen).values(type=record.type,
                                                           time=record.time,
                                                           push_time=record.push_time,
                                                           content=record.content,
                                                           sensitive=record.sensitive,
                                                           sensitive_score=record.sensitive_score,
                                                           url=record.url,
                                                           topic=record.topic,
                                                           username=username,
                                                           uuid=uuid))

                temp_human_llm_db.execute(insert(TempHumanLLM).values(type=record.type,
                                                                      time=record.time,
                                                                      push_time=record.push_time,
                                                                      content=record.content,
                                                                      is_bot=record.is_bot,
                                                                      is_bot_score=record.is_bot_score,
                                                                      model_judgment=record.model_judgment,
                                                                      url=record.url,
                                                                      topic=record.topic,
                                                                      username=username,
                                                                      uuid=uuid
                                                                      ))

                data_db.execute(insert(Data).values(type=record.type,
                                                    time=record.time,
                                                    push_time=record.time,
                                                    content=record.content,
                                                    sensitive=record.sensitive,
                                                    sensitive_score=record.sensitive_score,
                                                    is_bot=record.is_bot,
                                                    is_bot_score=record.is_bot_score,
                                                    model_judgment=record.model_judgment,
                                                    url=record.url,
                                                    topic=record.topic,
                                                    uuid=uuid
                                                    ))

                record.submitted = True

            daily_data_db.commit()
            temp_sen_db.commit()
            temp_human_llm_db.commit()
            data_db.commit()

        return JSONResponse(content={'message': f'{total_count} Records Submitted'})

    except Exception as e:
        return JSONResponse(status_code=status.HTTP_500_INTERNAL_SERVER_ERROR, content={'message': str(e)})


@app.put('/update/sen')
def update_sen(ID: str, approve: str, auditor: str,
               temp_sen_db: Session = Depends(get_db(SessionLocalTempSen)),
               sen_db: Session = Depends(get_db(SessionLocalSen)),
               data_db: Session = Depends(get_db(SessionLocalData))) -> JSONResponse:
    try:
        IDs = list(map(int, ID.split('&')))
        approves = list(map(int, approve.split('&')))
    except Exception as e:
        return JSONResponse(status_code=status.HTTP_400_BAD_REQUEST, content={'message': str(e)})
    try:
        submitted_time = datetime.now().strftime('%Y-%m-%d')
        for ID_, approve_ in zip(IDs, approves):
            if approve_:
                record = temp_sen_db.query(TempSen).filter(TempSen.ID == ID_).first()
                if not record:
                    continue

                sen_db.execute(insert(Sen).values(type=record.type,
                                                  time=record.time,
                                                  push_time=record.push_time,
                                                  submitted_time=submitted_time,
                                                  content=record.content,
                                                  sensitive=record.sensitive,
                                                  sensitive_score=record.sensitive_score,
                                                  url=record.url,
                                                  topic=record.topic,
                                                  auditor=auditor,
                                                  uuid=record.uuid
                                                  ))
                sen_db.commit()

                data_record = data_db.query(Data).filter(Data.uuid == record.uuid).first()
                if not data_record:
                    continue
                data_record.sensitive = record.sensitive
                data_record.sensitive_score = record.sensitive_score
                data_db.commit()

            temp_sen_db.execute(delete(TempSen).filter(TempSen.ID == ID_))
            temp_sen_db.commit()

        return JSONResponse(content={'message': 'Updated Successfully'})

    except Exception as e:
        import traceback
        error = traceback.format_exc()
        print(error)
        return JSONResponse(status_code=status.HTTP_500_INTERNAL_SERVER_ERROR, content={'message': str(e)})


@app.put('/update/human_llm')
def update_human_llm(ID: str, approve: str, auditor: str,
                     temp_human_llm_db: Session = Depends(get_db(SessionLocalTempHumanLLM)),
                     human_llm_db: Session = Depends(get_db(SessionLocalHumanLLM)),
                     data_db: Session = Depends(get_db(SessionLocalData))) -> JSONResponse:
    try:
        IDs = list(map(int, ID.split('&')))
        approves = list(map(int, approve.split('&')))
    except Exception as e:
        return JSONResponse(status_code=status.HTTP_400_BAD_REQUEST, content={'message': str(e)})
    try:
        submitted_time = datetime.now().strftime('%Y-%m-%d')
        for ID_, approve_ in zip(IDs, approves):
            if approve_:
                record = temp_human_llm_db.query(TempHumanLLM).filter(TempHumanLLM.ID == ID_).first()
                if not record:
                    continue

                human_llm_db.execute(insert(HumanLLM).values(type=record.type,
                                                             time=record.time,
                                                             push_time=record.push_time,
                                                             submitted_time=submitted_time,
                                                             content=record.content,
                                                             is_bot=record.is_bot,
                                                             is_bot_score=record.is_bot_score,
                                                             model_judgment=record.model_judgment,
                                                             url=record.url,
                                                             topic=record.topic,
                                                             auditor=auditor,
                                                             uuid=record.uuid
                                                             ))
                human_llm_db.commit()

                data_record = data_db.query(Data).filter(Data.uuid == record.uuid).first()
                if data_record is not None:
                    data_record.is_bot = record.is_bot
                    data_record.is_bot_score = record.is_bot_score
                    data_record.model_judgment = record.model_judgment
                    data_db.commit()

            temp_human_llm_db.execute(delete(TempHumanLLM).filter(TempHumanLLM.ID == ID_))
            temp_human_llm_db.commit()

        return JSONResponse(content={'message': 'Updated Successfully'})

    except Exception as e:
        import traceback
        error = traceback.format_exc()
        print(error)
        return JSONResponse(status_code=status.HTTP_500_INTERNAL_SERVER_ERROR, content={'message': str(e)})


@app.delete('/delete/sen/{ID}')
def delete_sen(ID: str, db: Session = Depends(get_db(SessionLocalSen))) -> JSONResponse:
    try:
        IDs = list(map(int, ID.split('&')))
    except Exception as e:
        return JSONResponse(status_code=status.HTTP_400_BAD_REQUEST, content={'message': 'Bad Request'})
    try:
        for ID_ in IDs:
            db.execute(delete(Sen).filter(Sen.ID == ID_))
            db.commit()

        return JSONResponse(content={'message': 'Record Deleted'})
    except Exception as e:
        return JSONResponse(status_code=status.HTTP_500_INTERNAL_SERVER_ERROR, content={'message': str(e)})


@app.delete('/delete/human_llm/{ID}')
def delete_human_llm(ID: str, db: Session = Depends(get_db(SessionLocalHumanLLM))) -> JSONResponse:
    try:
        IDs = list(map(int, ID.split('&')))
    except Exception as e:
        return JSONResponse(status_code=status.HTTP_400_BAD_REQUEST, content={'message': 'Bad Request'})
    try:
        for ID_ in IDs:
            db.execute(delete(HumanLLM).filter(HumanLLM.ID == ID_))
            db.commit()

        return JSONResponse(content={'message': 'Record Deleted'})
    except Exception as e:
        return JSONResponse(status_code=status.HTTP_500_INTERNAL_SERVER_ERROR, content={'message': str(e)})


@app.get('/read_temp_sen')
def read_temp_sen(page: int = 1, page_size: int = 10, type: str = None, sensitive: bool = None,
                  content_keyword: str = None, topic: str = None, username: str = None,
                  db: Session = Depends(get_db(SessionLocalTempSen))) -> JSONResponse:
    try:
        filters = []
        if type:
            filters.append(TempSen.type == type)
        if sensitive is not None:
            filters.append(TempSen.sensitive == sensitive)
        if content_keyword:
            filters.append(TempSen.content.like(f'%{content_keyword}%'))
        if topic:
            filters.append(TempSen.topic.like(f'%{topic}%'))
        if username:
            filters.append(TempSen.username == username)

        query = db.query(TempSen).filter(*filters)
        total_count = query.count()

        results = query.limit(page_size).offset((page - 1) * page_size).all()

        rows = []
        for row in results:
            rows.append(
                {'ID': row.ID, 'type': row.type, 'time': row.time, 'content': row.content, 'sensitive': row.sensitive,
                 'sensitive_score': round(row.sensitive_score, 2), 'url': row.url, 'topic': row.topic,
                 'username': row.username})

        return JSONResponse(content={'total_count': total_count, 'temp_sen': rows})
    except Exception as e:
        return JSONResponse(status_code=status.HTTP_400_BAD_REQUEST, content={'message': str(e)})


@app.get('/read_temp_human_llm')
def read_human_llm(page: int = 1, page_size: int = 10, type: str = None, is_bot: bool = None,
                   model_judgment: str = None, content_keyword: str = None, topic: str = None,
                   username: str = None,
                   db: Session = Depends(get_db(SessionLocalTempHumanLLM))) -> JSONResponse:
    try:
        filters = []
        if type:
            filters.append(TempHumanLLM.type == type)
        if is_bot is not None:
            filters.append(TempHumanLLM.is_bot == is_bot)
        if model_judgment:
            filters.append(TempHumanLLM.model_judgment == model_judgment)
        if content_keyword:
            filters.append(TempHumanLLM.content.like(f'%{content_keyword}%'))
        if topic:
            filters.append(TempHumanLLM.topic.like(f'%{topic}%'))
        if username:
            filters.append(TempHumanLLM.username == username)

        query = db.query(TempHumanLLM).filter(*filters)
        total_count = query.count()

        results = query.limit(page_size).offset((page - 1) * page_size).all()

        rows = []
        for row in results:
            rows.append({'ID': row.ID, 'type': row.type, 'time': row.time, 'content': row.content, 'is_bot': row.is_bot,
                         'is_bot_score': round(row.is_bot_score, 2), 'model_judgment': row.model_judgment,
                         'url': row.url,
                         'topic': row.topic, 'username': row.username})

        return JSONResponse(content={'total_count': total_count, 'temp_human_llm': rows})
    except Exception as e:
        return JSONResponse(status_code=status.HTTP_400_BAD_REQUEST, content={'message': str(e)})


@app.post('/login')
def login(username: str, password: str,
          db: Session = Depends(get_db(SessionLocalUsers))) -> JSONResponse:
    try:
        record = db.query(Users.password, Users.role).filter(Users.username == username).first()
        if record is None:
            return JSONResponse(status_code=status.HTTP_404_NOT_FOUND, content={'message': 'User Not Found'})
        correct_password = record.password
        role = record.role
        if password == correct_password:
            return JSONResponse(content={'message': role})
        else:
            return JSONResponse(status_code=status.HTTP_401_UNAUTHORIZED, content={'message': 'Wrong Password'})
    except Exception as e:
        return JSONResponse(status_code=status.HTTP_500_INTERNAL_SERVER_ERROR, content={'message': str(e)})


@app.post('/register')
def register(username: str, password: str,
             db: Session = Depends(get_db(SessionLocalUsers))) -> JSONResponse:
    try:
        db.execute(insert(Users).values(username=username,
                                        password=password,
                                        role='user'))
        db.commit()

        return JSONResponse(status_code=status.HTTP_201_CREATED, content={'message': 'User Registered Successfully'})
    except IntegrityError:
        return JSONResponse(status_code=status.HTTP_409_CONFLICT, content={'message': 'Username Already Exists'})
    except Exception as e:
        return JSONResponse(status_code=status.HTTP_500_INTERNAL_SERVER_ERROR, content={'message': str(e)})


@app.put('/change_password/{username}')
def change_password(username: str, old_password: str, new_password: str,
                    db: Session = Depends(get_db(SessionLocalUsers))) -> JSONResponse:
    try:
        record = db.query(Users.password).filter(Users.username == username).first()

        if record is None:
            return JSONResponse(status_code=status.HTTP_404_NOT_FOUND, content={'message': 'User Not Found'})
        correct_password = record.password
        if old_password == correct_password:
            record.password = new_password
            db.commit()

            return JSONResponse(content={'message': 'Password Changed Successfully'})
        else:
            return JSONResponse(status_code=status.HTTP_401_UNAUTHORIZED, content={'message': 'Wrong Old Password'})
    except Exception as e:
        return JSONResponse(status_code=status.HTTP_500_INTERNAL_SERVER_ERROR, content={'message': str(e)})


@app.get('/search_account')
def search_account(page: int = 1, page_size: int = 10, username: str = None, role: str = None,
                   db: Session = Depends(get_db(SessionLocalUsers))) -> JSONResponse:
    try:
        filters = []
        if username:
            filters.append(Users.username == username)
        if role:
            filters.append(Users.role == role)

        query = db.query(Users).filter(*filters)
        total_count = query.count()

        results = query.limit(page_size).offset((page - 1) * page_size).all()

        rows = []
        for row in results:
            rows.append({'ID': row.ID, 'username': row.username, 'password': row.password, 'role': row.role})

        return JSONResponse(content={'total_count': total_count, 'accounts': rows})
    except Exception as e:
        return JSONResponse(status_code=status.HTTP_500_INTERNAL_SERVER_ERROR, content={'message': str(e)})


@app.put('/update_account/{ID}')
def update(ID: int, username: str = None, password: str = None, role: str = None,
           db: Session = Depends(get_db(SessionLocalUsers))) -> JSONResponse:
    try:
        record = db.query(Users).filter(Users.ID == ID).first()
        if username:
            record.username = username
        if password:
            record.password = password
        if role:
            record.role = role
        db.commit()

        return JSONResponse(content={'message': 'Account Updated Successfully'})
    except IntegrityError:
        return JSONResponse(status_code=status.HTTP_409_CONFLICT, content={'message': 'Username Already Exists'})
    except Exception as e:
        return JSONResponse(status_code=status.HTTP_500_INTERNAL_SERVER_ERROR, content={'message': str(e)})


@app.delete('/delete/account/{ID}')
def delete_account(ID: str,
                   db: Session = Depends(get_db(SessionLocalUsers))) -> JSONResponse:
    try:
        IDs = list(map(int, ID.split('&')))
    except Exception as e:
        return JSONResponse(status_code=status.HTTP_400_BAD_REQUEST, content={'message': 'Bad Request'})
    try:
        for ID_ in IDs:
            result = db.execute(delete(Users).filter(Users.ID == ID_))
            db.commit()

        return JSONResponse(content={'message': 'Account Deleted'})
    except Exception as e:
        return JSONResponse(status_code=status.HTTP_500_INTERNAL_SERVER_ERROR, content={'message': str(e)})


@app.get('/analysis_report')
def analysis_report(start_date: str, end_date: str, content_keyword: str, width: int = 800, height: int = 400,
                    background_color: str = 'white') -> JSONResponse:
    start_date = datetime.strptime(start_date, '%Y-%m-%d')
    end_date = datetime.strptime(end_date, '%Y-%m-%d')
    base_query = '''
        SELECT type, strftime('%Y-%m-%d', time), content, {keys}
        FROM {database}
        WHERE time IS NOT NULL AND time BETWEEN :sd AND :ed AND topic LIKE :t
    '''
    params = {'sd': start_date, 'ed': end_date, 't': f'%{content_keyword}%'}
    try:
        results_container = {}

        def execute_query(session_local, query, params, key):
            db = session_local()
            results = db.execute(text(query), params).fetchall()
            rows = []
            for row in results:
                rows.append(list(row))

            results_container[key] = rows
            db.close()

        sen_thread = threading.Thread(target=execute_query,
                                      args=(SessionLocalSen,
                                            base_query.format(keys='sensitive', database='sen'),
                                            params,
                                            'sen'))
        human_thread = threading.Thread(target=execute_query,
                                        args=(SessionLocalHumanLLM,
                                              base_query.format(keys='is_bot', database='human_llm'),
                                              params,
                                              'human_llm'))

        sen_thread.start()
        human_thread.start()
        sen_thread.join()
        human_thread.join()

        sen_data = results_container['sen']
        human_llm_data = results_container['human_llm']

        combined_content = set()

        sen_distribution_dict = {}
        for type in type_list:
            sen_distribution_dict[type] = {'count': [0, 0], 'date': {}}

        sen_total_count = 0
        for record in sen_data:
            type, date, content, is_sensitive = record
            if date is None:
                continue
            sen_total_count += int(is_sensitive)
            if is_sensitive:
                combined_content.add(content)
            sen_distribution_dict[type]['count'][0] += int(is_sensitive)
            sen_distribution_dict[type]['count'][1] += 1
            if date not in sen_distribution_dict[type]['date']:
                sen_distribution_dict[type]['date'][date] = [0, 0]
            sen_distribution_dict[type]['date'][date][0] += int(is_sensitive)
            sen_distribution_dict[type]['date'][date][1] += 1

        human_distribution_dict = {}
        for type in type_list:
            human_distribution_dict[type] = {'count': [0, 0], 'date': {}}

        human_total_count = 0
        for record in human_llm_data:
            type, date, content, is_bot = record
            if date is None:
                continue
            human_total_count += int(is_bot)
            if is_bot:
                combined_content.add(content)
            human_distribution_dict[type]['count'][0] += int(is_bot)
            human_distribution_dict[type]['count'][1] += 1
            if date not in human_distribution_dict[type]['date']:
                human_distribution_dict[type]['date'][date] = [0, 0]
            human_distribution_dict[type]['date'][date][0] += int(is_bot)
            human_distribution_dict[type]['date'][date][1] += 1

        if len(combined_content) == 0:
            return JSONResponse(status_code=status.HTTP_404_NOT_FOUND, content={'message': 'No Record'})

        if sen_total_count == 0:
            sen_total_count = 1
        if human_total_count == 0:
            human_total_count = 1
        for type in type_list:
            if sen_distribution_dict[type]['count'][1] == 0:
                sen_distribution_dict[type]['count'][1] = 1
            if human_distribution_dict[type]['count'][1] == 0:
                human_distribution_dict[type]['count'][1] = 1

        if len(combined_content) > 1000:
            combined_content = random.sample(list(combined_content), 1000)

        combined_text = ' '.join(combined_content)
        filtered_combined_text = [word for word in jieba.cut(combined_text) if word not in stopwords and word.strip()]
        filtered_combined_text = ' '.join(filtered_combined_text)
        keywords = jieba.analyse.extract_tags(filtered_combined_text, topK=10)

        report = f"针对重点话题{content_keyword}，其在{start_date.strftime('%Y年%#m月%#d日')}至{end_date.strftime('%Y年%m月%d日')}的关键词如下：\n"
        report += '、'.join(keywords) + '。\n'

        report += "对其进行涉华敏感性分析的结果如下：\n敏感分布如下：\n"
        for type in type_list:
            report += f"在{type}中检测到的涉华敏感信息占比为{sen_distribution_dict[type]['count'][0] / sen_distribution_dict[type]['count'][1] * 100:.1f}%，"
        report = report[:-1] + '。\n' + '涉华敏感信息的来源分布如下：'
        for type in type_list:
            report += f"{sen_distribution_dict[type]['count'][0] / sen_total_count * 100:.1f}%的敏感信息来自{type}，"
        report = report[:-1] + '。\n' + '从时间分布上讲，'
        for type in type_list:
            try:
                peak_date = max(sen_distribution_dict[type]['date'],
                                key=sen_distribution_dict[type]['date'].get)
                report += f"{type}中的信息在{datetime.strptime(peak_date, '%Y-%m-%d').strftime('%Y年%#m月%#d日')}上涉华敏感信息的数量达到峰值，"
            except:
                report += f"{type}中的信息暂无时间分布，"
        report = report[:-1] + '。\n' + '其传播趋势如下：'
        current_date = datetime.now()
        one_month_ago = current_date - timedelta(days=30)
        for type in type_list:
            try:
                peak_date = max(sen_distribution_dict[type]['date'],
                                key=sen_distribution_dict[type]['date'].get)
                peak_date_dt = datetime.strptime(peak_date, '%Y-%m-%d')
                if peak_date_dt >= one_month_ago:
                    report += f"{type}中该话题仍处于强敏感环境中，潜在传播风险比较高，需要对其进行重点观测和防备。"
                else:
                    report += f"{type}中该话题最近处于低敏感环境中，潜在传播风险比较低。"
            except:
                report += f"{type}中该话题最近处于低敏感环境中，潜在传播风险比较低。"

        report = report[:-1] + '。\n' + '生成式的分布如下：'
        for type in type_list:
            report += f"在{type}中检测到的为大语言模型生成的信息占比为{human_distribution_dict[type]['count'][0] / human_distribution_dict[type]['count'][1] * 100:.1f}%，"
        report = report[:-1] + '。\n' + '生成式信息的来源分布如下：'
        for type in type_list:
            report += f"{human_distribution_dict[type]['count'][0] / human_total_count * 100:.1f}%的生成式信息来自{type}，"
        report = report[:-1] + '。\n' + '从时间分布上讲，'
        for type in type_list:
            try:
                peak_date = max(human_distribution_dict[type]['date'],
                                key=human_distribution_dict[type]['date'].get)
                report += f"{type}中的信息在{datetime.strptime(peak_date, '%Y-%m-%d').strftime('%Y年%#m月%#d日')}上生成式信息的数量达到峰值，"
            except:
                report += f"{type}中的信息暂无时间分布，"
        report = report[:-1] + '。\n'

        sen_type_distribution = {}
        for type in type_list:
            sen_type_distribution[type] = sen_distribution_dict[type]['count'][0]

        human_llm_type_distribution = {}
        for type in type_list:
            human_llm_type_distribution[type] = human_distribution_dict[type]['count'][0]

        sen_time_distribution = {}
        for type in type_list:
            for date in sen_distribution_dict[type]['date']:
                if date not in sen_time_distribution:
                    sen_time_distribution[date] = {}
                sen_time_distribution[date][type] = {
                    '1': sen_distribution_dict[type]['date'][date][0],
                    '0': sen_distribution_dict[type]['date'][date][1] - sen_distribution_dict[type]['date'][date][0]
                }

        human_llm_time_distribution = {}
        for type in type_list:
            for date in human_distribution_dict[type]['date']:
                if date not in human_llm_time_distribution:
                    human_llm_time_distribution[date] = {}
                human_llm_time_distribution[date][type] = {
                    '1': human_distribution_dict[type]['date'][date][0],
                    '0': human_distribution_dict[type]['date'][date][1] - human_distribution_dict[type]['date'][date][0]
                }

        wc = WordCloud(width=width, height=height, background_color=background_color, stopwords=STOPWORDS,
                       font_path='./utils/simhei.ttf').generate(filtered_combined_text)
        img_buffer = io.BytesIO()
        wc.to_image().save(img_buffer, format='PNG')
        img_buffer.seek(0)
        img_data = img_buffer.read()

        report_id = uuid4().hex
        report_data_store[report_id] = {
            'topic': content_keyword,
            'start_date': start_date,
            'end_date': end_date,
            'report': report,
            'keywords': keywords,
            'wordcloud': base64.b64encode(img_data).decode('utf-8'),
            'timestamp': datetime.now()
        }

    except Exception as e:
        return JSONResponse(status_code=status.HTTP_500_INTERNAL_SERVER_ERROR, content={'message': str(e)})

    return JSONResponse(content={
        'report': report,
        'keywords': keywords,
        'sen_type_distribution': sen_type_distribution,
        'human_llm_type_distribution': human_llm_type_distribution,
        'sen_time_distribution': dict(sorted(sen_time_distribution.items(), key=lambda item: item[0])),
        'human_llm_time_distribution': dict(sorted(human_llm_time_distribution.items(), key=lambda item: item[0])),
        'wordcloud': base64.b64encode(img_data).decode('utf-8'),
        'report_id': report_id})


@app.get('/generate_report')
def generate_report(report_id: str):
    if report_id not in report_data_store:
        return JSONResponse(status_code=status.HTTP_404_NOT_FOUND, content={'message': 'Report Not Found'})

    try:
        data = report_data_store[report_id]

        document = Document()
        document.styles['Normal'].font.name = u'宋体'
        document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

        head = document.add_heading('', level=1)
        run = head.add_run(f"多维特征分析报告：{data['topic']}")
        run.font.name = u'宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), u'Cambria')

        document.add_paragraph(f"时间：{data['start_date']} 至 {data['end_date']}")
        document.add_paragraph(f"关键词：{'、'.join(data['keywords'])}。")

        head = document.add_heading('', level=2)
        run = head.add_run("词云图")
        run.font.name = u'宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), u'Cambria')

        image_stream = io.BytesIO(base64.b64decode(data['wordcloud']))
        picture = document.add_picture(image_stream)
        picture.width = Cm(15)
        picture.height = Cm(7.5)

        head = document.add_heading('', level=2)
        run = head.add_run("文字报告内容")
        run.font.name = u'宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), u'Cambria')

        for p in data['report'].split('\n'):
            document.add_paragraph(f"   {p}")

        file_stream = io.BytesIO()
        document.save(file_stream)
        file_stream.seek(0)

        return StreamingResponse(file_stream,
                                 media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                 headers={"Content-Disposition": f"attachment;filename={report_id}.docx"})
    except Exception as e:
        return JSONResponse(status_code=status.HTTP_500_INTERNAL_SERVER_ERROR, content={'message': str(e)})


@app.middleware('http')
async def process_time_middleware(request: Request, call_next):
    t0 = time.time()
    query_params = request.query_params._dict
    cleaned_params = {}
    for k, v in query_params.items():
        if v != '':
            cleaned_params[k] = v

    request.scope['query_string'] = urlencode(cleaned_params, doseq=True).encode('utf-8')
    response = await call_next(request)
    print('Request: {}  {}'.format(time.strftime('%Y-%m-%d %H:%M:%S', time.localtime(t0)), time.time() - t0))
    return response


def clear_expired_reports():
    now = datetime.now()
    expired_ids = [report_id for report_id, data in report_data_store.items() if
                   now - data['timestamp'] > timedelta(minutes=5)]
    for report_id in expired_ids:
        del report_data_store[report_id]


def clear_cached_items():
    now = datetime.now()
    expired_keys = [key for key, expire_time in cache_expiry_times.items() if
                    now > expire_time]
    for expired_key in expired_keys:
        del cached_item_store[expired_key]
        del cache_expiry_times[expired_key]


@app.on_event('startup')
async def startup_event():
    schedular.add_job(clear_expired_reports, IntervalTrigger(minutes=1))
    schedular.add_job(clear_cached_items, IntervalTrigger(seconds=10))
    schedular.start()


if __name__ == '__main__':
    init_db()
    uvicorn.run(app, host='0.0.0.0', port=8123, log_level='debug')
